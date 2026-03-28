import re

from docx import Document

from app.ingestion.pdf import (
    CONDITIONAL_PATTERN,
    NUMBERED_STEP_PATTERN,
    ParsedDocument,
    Section,
)


def parse_docx(file_bytes: bytes, filename: str = "") -> ParsedDocument:
    from io import BytesIO

    doc = Document(BytesIO(file_bytes))
    sections: list[Section] = []
    raw_parts: list[str] = []
    current_section: Section | None = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()
        is_heading = style_name.startswith("heading")

        if is_heading:
            if current_section:
                sections.append(current_section)
            level = int(style_name.replace("heading", "").strip() or "1")
            current_section = Section(
                header=text, content="", level=level, markers={}
            )
        else:
            if current_section is None:
                current_section = Section(
                    header="", content="", level=0, markers={}
                )
            current_section.content += text + "\n"

            if NUMBERED_STEP_PATTERN.match(text):
                current_section.markers["has_numbered_steps"] = True
            if CONDITIONAL_PATTERN.search(text):
                current_section.markers["has_conditionals"] = True
            if "list" in style_name:
                current_section.markers["has_list"] = True

        raw_parts.append(text)

    if current_section:
        sections.append(current_section)

    raw_text = "\n".join(raw_parts)

    return ParsedDocument(
        title=filename or "Untitled DOCX",
        sections=sections,
        raw_text=raw_text,
        metadata={"paragraph_count": len(doc.paragraphs), "source_type": "docx"},
    )
